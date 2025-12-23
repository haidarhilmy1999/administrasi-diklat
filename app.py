import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import datetime
import zipfile
import time

# --- LIBRARY GOOGLE SHEETS ---
import gspread
from oauth2client.service_account import ServiceAccountCredentials

# =============================================================================
# 1. KONFIGURASI HALAMAN
# =============================================================================
st.set_page_config(
    page_title="Admin Diklat BC", 
    layout="wide", 
    page_icon="⚡",
    initial_sidebar_state="collapsed" 
)

# CSS STYLING
hide_st_style = """
            <style>
            #MainMenu {visibility: hidden;}
            footer {visibility: hidden;}
            header {visibility: hidden;} 
            [data-testid="stToolbar"] {visibility: hidden;}
            [data-testid="stDecoration"] {display: none;}
            .stAppDeployButton {display: none !important;}
            div[class*="viewerBadge"] {display: none !important;}
            .block-container {padding-top: 1rem;}
            
            .danger-box {
                border: 1px solid #ff4b4b;
                padding: 10px;
                border-radius: 5px;
                background-color: #fff5f5;
                color: #ff4b4b;
            }
            </style>
            """
st.markdown(hide_st_style, unsafe_allow_html=True)

# KONFIGURASI DATABASE
NAMA_GOOGLE_SHEET = "Database_Diklat_DJBC"
SHEET_HISTORY = "Sheet1"
SHEET_KALENDER = "Master_Kalender"

if 'history_log' not in st.session_state:
    st.session_state['history_log'] = pd.DataFrame(columns=['TIMESTAMP', 'NAMA', 'NIP', 'DIKLAT', 'SATKER'])
if 'uploader_key' not in st.session_state:
    st.session_state['uploader_key'] = 0

# --- KONEKSI DATABASE ---
def connect_to_gsheet():
    try:
        scope = ['https://spreadsheets.google.com/feeds', 'https://www.googleapis.com/auth/drive']
        creds_dict = dict(st.secrets["gcp_service_account"])
        creds = ServiceAccountCredentials.from_json_keyfile_dict(creds_dict, scope)
        client = gspread.authorize(creds)
        return client.open(NAMA_GOOGLE_SHEET)
    except: return None

# =============================================================================
# 2. LOGIKA DATABASE
# =============================================================================

def format_date_range(row):
    try:
        tgl_mulai = row['TANGGAL_MULAI']
        tgl_selesai = row['TANGGAL_SELESAI']
        if isinstance(tgl_mulai, (pd.Timestamp, datetime.datetime)): tgl_mulai = tgl_mulai.strftime("%d %b %Y")
        if isinstance(tgl_selesai, (pd.Timestamp, datetime.datetime)): tgl_selesai = tgl_selesai.strftime("%d %b %Y")
        str_mulai = str(tgl_mulai).strip(); str_selesai = str(tgl_selesai).strip()
        return str_mulai if str_mulai == str_selesai else f"{str_mulai} s.d. {str_selesai}"
    except: return "-"

def update_calendar_db(df_new):
    sh = connect_to_gsheet()
    if sh:
        try:
            ws = sh.worksheet(SHEET_KALENDER)
            existing_data = ws.get_all_records()
            df_old = pd.DataFrame(existing_data)
            
            required = ['TANGGAL_MULAI', 'TANGGAL_SELESAI', 'LOKASI', 'JUDUL_PELATIHAN']
            if not all(col in df_new.columns for col in required):
                st.error(f"Excel harus punya kolom: {', '.join(required)}")
                return False

            df_new['RENCANA_TANGGAL'] = df_new.apply(format_date_range, axis=1)
            df_new = df_new.astype(str)
            if not df_old.empty: df_old = df_old.astype(str)

            final_rows = []
            last_id = 0
            if not df_old.empty and 'ID' in df_old.columns:
                try:
                    numeric_ids = pd.to_numeric(df_old['ID'], errors='coerce').fillna(0)
                    last_id = int(numeric_ids.max())
                except: pass
            
            processed_titles = []
            for _, row_new in df_new.iterrows():
                judul = row_new['JUDUL_PELATIHAN']
                processed_titles.append(judul)
                match_old = df_old[df_old['JUDUL_PELATIHAN'] == judul] if not df_old.empty else pd.DataFrame()
                
                if not match_old.empty:
                    row_lama = match_old.iloc[0]
                    final_rows.append([
                        row_lama['ID'], judul, row_new['RENCANA_TANGGAL'], row_new['LOKASI'], 
                        row_lama['STATUS'], row_lama['REALISASI']
                    ])
                else:
                    last_id += 1
                    final_rows.append([last_id, judul, row_new['RENCANA_TANGGAL'], row_new['LOKASI'], "Pending", "-"])
            
            if not df_old.empty:
                sisa_lama = df_old[~df_old['JUDUL_PELATIHAN'].isin(processed_titles)]
                for _, row_old in sisa_lama.iterrows():
                    final_rows.append(row_old.tolist())

            ws.clear()
            ws.append_row(['ID', 'JUDUL_PELATIHAN', 'RENCANA_TANGGAL', 'LOKASI', 'STATUS', 'REALISASI'])
            ws.append_rows(final_rows)
            st.toast("✅ Kalender berhasil di-update!", icon="📅")
            return True
        except Exception as e:
            st.error(f"Gagal update: {e}")
            return False
    return False

def mark_training_complete(judul_pelatihan):
    sh = connect_to_gsheet()
    if sh:
        try:
            ws = sh.worksheet(SHEET_KALENDER)
            data = ws.get_all_records()
            df = pd.DataFrame(data)
            row_idx = df.index[df['JUDUL_PELATIHAN'] == judul_pelatihan].tolist()
            if row_idx:
                idx_gsheet = row_idx[0] + 2 
                current_time = datetime.datetime.now().strftime("%d-%m-%Y")
                ws.update_cell(idx_gsheet, 5, "Selesai") 
                ws.update_cell(idx_gsheet, 6, current_time) 
                return True
        except: pass
    return False

def save_to_cloud_callback(df_input):
    current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    try:
        data_to_save = df_input.copy()
        target_cols = ['NAMA', 'NIP', 'JUDUL_PELATIHAN', 'SATKER']
        final_cols = {}
        for target in target_cols:
            for col in data_to_save.columns:
                if target in col: 
                    final_cols[col] = target if target != 'JUDUL_PELATIHAN' else 'DIKLAT'
                    break
        if final_cols:
            data_to_save = data_to_save[list(final_cols.keys())].rename(columns=final_cols)
            data_to_save.insert(0, 'TIMESTAMP', current_time)
            
            sh = connect_to_gsheet()
            if sh:
                ws_log = sh.worksheet(SHEET_HISTORY)
                ws_log.append_rows(data_to_save.astype(str).values.tolist())
                st.toast("✅ Log tersimpan!", icon="☁️")
                
                if 'DIKLAT' in data_to_save.columns:
                    unique_titles = data_to_save['DIKLAT'].unique()
                    count = 0
                    for judul in unique_titles:
                        if mark_training_complete(judul): count += 1
                    if count > 0: st.toast(f"✅ {count} Jadwal Kalender ditandai Selesai!", icon="🎯")
            else: st.toast("⚠️ Gagal koneksi Cloud.", icon="📂")
    except Exception as e: st.toast(f"Error Database: {e}", icon="❌")

def reset_calendar_status():
    sh = connect_to_gsheet()
    if sh:
        try:
            ws = sh.worksheet(SHEET_KALENDER)
            all_values = ws.get_all_values()
            if len(all_values) > 1: 
                new_data = []
                for row in all_values[1:]:
                    while len(row) < 6: row.append("")
                    row[4] = "Pending"
                    row[5] = "-"
                    new_data.append(row)
                range_update = f"A2:F{len(all_values)}"
                ws.update(range_name=range_update, values=new_data)
                st.success("✅ Status Kalender berhasil di-reset menjadi 'Pending'!")
            else: st.warning("Data kalender kosong.")
        except Exception as e: st.error(f"Gagal reset kalender: {e}")

def clear_history_log():
    sh = connect_to_gsheet()
    if sh:
        try:
            ws = sh.worksheet(SHEET_HISTORY)
            ws.clear()
            ws.append_row(['TIMESTAMP', 'NAMA', 'NIP', 'DIKLAT', 'SATKER'])
            st.success("✅ Log Peserta berhasil dikosongkan!")
        except Exception as e: st.error(f"Gagal hapus log: {e}")

def reset_app():
    st.session_state['uploader_key'] += 1
    st.rerun()

# =============================================================================
# 3. FUNGSI UTILS & WORD GENERATOR (A4 + FIX HEADER & NUMBERING)
# =============================================================================
def calculate_age_from_nip(nip_str):
    try:
        clean_nip = str(nip_str).replace(" ", "").replace(".", "").replace("-", "")
        year_str = clean_nip[:4]
        if year_str.isdigit():
            birth_year = int(year_str); current_year = datetime.datetime.now().year
            if 1950 <= birth_year <= current_year: return current_year - birth_year
        return None
    except: return None

def get_gender_from_nip(nip_str):
    try:
        clean_nip = str(nip_str).replace(" ", "").replace(".", "").replace("-", "")
        if len(clean_nip) >= 15:
            code = clean_nip[14]; return "Pria" if code == '1' else "Wanita" if code == '2' else "Tidak Diketahui"
        return "Tidak Diketahui"
    except: return "Tidak Diketahui"

def set_repeat_table_header(row):
    tr = row._tr; trPr = tr.get_or_add_trPr(); tblHeader = OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'), "true"); trPr.append(tblHeader)

# --- FUNGSI WORD SINGLE (ZIP) ---
def create_single_document(row, judul, tgl_pel, tempat_pel, nama_ttd, jabatan_ttd, no_nd_val, tgl_nd_val):
    JENIS_FONT = 'Arial'; UKURAN_FONT = 11
    doc = Document(); style = doc.styles['Normal']; style.font.name = JENIS_FONT; style.font.size = Pt(UKURAN_FONT)
    section = doc.sections[0]
    section.page_width = Mm(210); section.page_height = Mm(297)
    section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0); section.right_margin = Cm(2.5)

    header_table = doc.add_table(rows=4, cols=3); header_table.alignment = WD_TABLE_ALIGNMENT.RIGHT 
    header_table.columns[0].width = Cm(2.0); header_table.columns[2].width = Cm(5.0)
    
    def isi_sel(r, c, text, size=11, bold=False):
        cell = header_table.cell(r, c); p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text); run.font.name = JENIS_FONT; run.font.size = Pt(size); run.bold = bold
    
    isi_sel(0, 0, "LAMPIRAN II"); header_table.cell(0, 2).merge(header_table.cell(0, 0))
    isi_sel(1, 0, f"Nota Dinas {jabatan_ttd}", size=11); header_table.cell(1, 2).merge(header_table.cell(1, 0))
    isi_sel(2, 0, "Nomor"); isi_sel(2, 1, ":"); isi_sel(2, 2, str(no_nd_val), size=8) # Ukuran 8
    isi_sel(3, 0, "Tanggal"); isi_sel(3, 1, ":"); isi_sel(3, 2, str(tgl_nd_val), size=8) # Ukuran 8
    
    doc.add_paragraph(""); p = doc.add_paragraph("DAFTAR PESERTA PELATIHAN"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].bold = True
    info_table = doc.add_table(rows=3, cols=3); 
    infos = [("Nama Pelatihan", judul), ("Tanggal", tgl_pel), ("Lokasi", tempat_pel)]
    for r, (l, v) in enumerate(infos): info_table.cell(r,0).text = l; info_table.cell(r,1).text = ":"; info_table.cell(r,2).text = str(v)
    
    doc.add_paragraph(""); table = doc.add_table(rows=2, cols=5); table.style = 'Table Grid'
    headers = ['NO', 'NAMA PEGAWAI', 'NIP', 'PANGKAT - GOL', 'SATUAN KERJA']; widths = [Cm(1.0), Cm(5.0), Cm(3.8), Cm(3.5), Cm(3.5)]
    for i in range(5): 
        table.rows[0].cells[i].text = headers[i]; table.rows[0].cells[i].width = widths[i]; table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    # Nomor selalu 1 karena single doc
    vals = ["1", row.get('NAMA','-'), row.get('NIP','-'), row.get('PANGKAT','-'), row.get('SATKER','-')]
    for i in range(5): table.rows[1].cells[i].text = str(vals[i])
    
    doc.add_paragraph("")
    ttd_table = doc.add_table(rows=1, cols=2); ttd_table.autofit = False
    ttd_table.columns[0].width = Cm(8.0); ttd_table.columns[1].width = Cm(7.5)
    
    p = ttd_table.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(f"{jabatan_ttd},"); p.add_run("\n\n\n\n\n\n") 
    run_elec = p.add_run("Ditandatangani secara elektronik")
    run_elec.font.size = Pt(10); run_elec.font.color.rgb = RGBColor(160, 160, 160)
    p.add_run(f"\n{nama_ttd}")
    
    f_out = io.BytesIO(); doc.save(f_out); f_out.seek(0); return f_out

# --- FUNGSI WORD COMBINED (FIXED HEADER & NUMBERING) ---
def generate_word_combined(df, nama_ttd, jabatan_ttd, no_nd_val, tgl_nd_val):
    JENIS_FONT = 'Arial'; UKURAN_FONT = 11; output = io.BytesIO(); doc = Document()
    style = doc.styles['Normal']; style.font.name = JENIS_FONT; style.font.size = Pt(UKURAN_FONT)
    
    section = doc.sections[0]
    section.page_width = Mm(210); section.page_height = Mm(297)
    section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0); section.right_margin = Cm(2.5)

    p_foot = section.footer.paragraphs[0]; p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER; run = p_foot.add_run(); run._r.append(OxmlElement('w:fldChar')); run._r[-1].set(qn('w:fldCharType'), 'begin'); run._r.append(OxmlElement('w:instrText')); run._r[-1].text = "PAGE"; run._r.append(OxmlElement('w:fldChar')); run._r[-1].set(qn('w:fldCharType'), 'end')
    
    # --- HEADER DISINI (SEKALI SAJA DI ATAS) ---
    header_table = doc.add_table(rows=4, cols=3); header_table.alignment = WD_TABLE_ALIGNMENT.RIGHT 
    header_table.columns[0].width = Cm(2.0); header_table.columns[2].width = Cm(5.0)
    def isi_sel(r, c, text, size=11, bold=False):
        cell = header_table.cell(r, c); p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text); run.font.name = JENIS_FONT; run.font.size = Pt(size); run.bold = bold
    
    isi_sel(0, 0, "LAMPIRAN II"); header_table.cell(0, 2).merge(header_table.cell(0, 0))
    isi_sel(1, 0, f"Nota Dinas {jabatan_ttd}", size=11); header_table.cell(1, 2).merge(header_table.cell(1, 0))
    isi_sel(2, 0, "Nomor"); isi_sel(2, 1, ":"); isi_sel(2, 2, str(no_nd_val), size=8) # FONT 8
    isi_sel(3, 0, "Tanggal"); isi_sel(3, 1, ":"); isi_sel(3, 2, str(tgl_nd_val), size=8) # FONT 8
    
    col_judul = 'JUDUL_PELATIHAN' if 'JUDUL_PELATIHAN' in df.columns else df.columns[0]
    kelompok = df.groupby(col_judul)
    total_groups = len(kelompok) 
    counter = 0 
    
    for judul, group in kelompok:
        counter += 1
        first = group.iloc[0]
        
        # Mulai Konten Per Pelatihan
        doc.add_paragraph("")
        p = doc.add_paragraph("DAFTAR PESERTA PELATIHAN"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].bold = True
        
        info_table = doc.add_table(rows=3, cols=3); info_table.autofit = False
        info_table.columns[0].width = Cm(4.0); info_table.columns[1].width = Cm(0.5); info_table.columns[2].width = Cm(11.0)
        infos = [("Nama Pelatihan", judul), ("Tanggal", first.get('TANGGAL_PELATIHAN','-')), ("Lokasi", first.get('TEMPAT','-'))]
        for r, (l, v) in enumerate(infos): info_table.cell(r,0).text = l; info_table.cell(r,1).text = ":"; info_table.cell(r,2).text = str(v)
        
        doc.add_paragraph("")
        table = doc.add_table(rows=1, cols=5); table.style = 'Table Grid'; table.autofit = False
        widths = [Cm(1.0), Cm(5.0), Cm(4.0), Cm(2.5), Cm(3.0)]
        
        hdr_cells = table.rows[0].cells; set_repeat_table_header(table.rows[0])
        headers = ['NO', 'NAMA PEGAWAI', 'NIP', 'PANGKAT', 'UNIT KERJA']
        for i in range(5): 
            hdr_cells[i].width = widths[i]
            p = hdr_cells[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(headers[i]); run.bold = True
            
        # --- RESET NOMOR URUT (ENUMERATE) ---
        for i, (idx, row) in enumerate(group.iterrows(), start=1):
            row_cells = table.add_row().cells
            vals = [str(i), row.get('NAMA','-'), row.get('NIP','-'), row.get('PANGKAT','-'), row.get('SATKER','-')]
            for k in range(5): 
                row_cells[k].width = widths[k]; row_cells[k].text = str(vals[k])
                row_cells[k].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        doc.add_paragraph("")
        
        # Tanda Tangan (Hanya di akhir)
        if counter == total_groups:
            ttd_table = doc.add_table(rows=1, cols=2); ttd_table.autofit = False
            ttd_table.columns[0].width = Cm(8.0); ttd_table.columns[1].width = Cm(7.5)
            
            p = ttd_table.cell(0, 1).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.add_run(f"{jabatan_ttd},"); p.add_run("\n\n\n\n\n\n") 
            run_elec = p.add_run("Ditandatangani secara elektronik")
            run_elec.font.size = Pt(10); run_elec.font.color.rgb = RGBColor(160, 160, 160)
            p.add_run(f"\n{nama_ttd}")
        
        if counter < total_groups: 
            doc.add_page_break()

    doc.save(output); output.seek(0); return output

def generate_zip_files(df, nama_ttd, jabatan_ttd, no_nd_val, tgl_nd_val):
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
        for idx, row in df.iterrows():
            judul = row.get('JUDUL_PELATIHAN', 'Diklat')
            nama_file = f"{str(row.get('NAMA','Peserta')).replace(' ', '_')}_{str(row.get('NIP','000'))}.docx"
            doc_buffer = create_single_document(row, judul, row.get('TANGGAL_PELATIHAN','-'), row.get('TEMPAT','-'), nama_ttd, jabatan_ttd, no_nd_val, tgl_nd_val)
            zip_file.writestr(nama_file, doc_buffer.getvalue())
    zip_buffer.seek(0); return zip_buffer

# =============================================================================
# 4. GUI UTAMA
# =============================================================================
st.title("Admin Diklat BC 🇮🇩")
st.markdown("---")

tab_gen, tab_cal, tab_dash, tab_db = st.tabs(["🚀 Generator", "📅 Kalender", "📊 Dashboard", "☁️ Database"])

# --- TAB KALENDER ---
with tab_cal:
    col_k1, col_k2 = st.columns([1, 2])
    with col_k1:
        st.subheader("Upload Kalender")
        st.info("Format: JUDUL | TANGGAL_MULAI | TANGGAL_SELESAI | LOKASI")
        file_kalender = st.file_uploader("Upload Excel Kalender", type=['xlsx'])
        
        # Template
        df_cal_dummy = pd.DataFrame({
            "JUDUL_PELATIHAN": ["DTSS Kepabeanan", "DTSD Cukai"],
            "TANGGAL_MULAI": ["12/01/2026", "02/02/2026"],
            "TANGGAL_SELESAI": ["16/01/2026", "05/02/2026"],
            "LOKASI": ["Pusdiklat BC", "KPU Batam"]
        })
        buf_cal = io.BytesIO(); 
        with pd.ExcelWriter(buf_cal, engine='xlsxwriter') as writer: df_cal_dummy.to_excel(writer, index=False)
        buf_cal.seek(0)
        st.download_button("📥 Template Kalender", buf_cal, "Template_Kalender.xlsx", use_container_width=True)

        if file_kalender:
            if st.button("Simpan / Update Kalender", type="primary"):
                df_new_cal = pd.read_excel(file_kalender)
                if 'JUDUL_PELATIHAN' in df_new_cal.columns:
                    update_calendar_db(df_new_cal)
                else: st.error("Format salah! Kolom JUDUL_PELATIHAN wajib ada.")

    with col_k2:
        st.subheader("Preview Master Kalender")
        sh = connect_to_gsheet()
        if sh:
            try:
                ws = sh.worksheet(SHEET_KALENDER)
                data_cal = ws.get_all_records()
                if data_cal: st.dataframe(pd.DataFrame(data_cal), use_container_width=True)
                else: st.warning("Data kalender masih kosong.")
            except: st.warning(f"Sheet '{SHEET_KALENDER}' belum dibuat di Google Sheets.")

# --- TAB GENERATOR (AUTO DETECT) ---
with tab_gen:
    c_up, c_ttd, c_nd = st.columns([1.5, 1.5, 1.5])
    with c_up:
        st.markdown("###### 1. Upload Data")
        uploaded_file = st.file_uploader("Upload Excel Peserta", type=['xlsx'], label_visibility="collapsed", key=f"uploader_{st.session_state['uploader_key']}")
        
        sc1, sc2 = st.columns(2)
        with sc1:
            df_dummy = pd.DataFrame({"JUDUL_PELATIHAN": ["DTSS Kepabeanan"], "TANGGAL_PELATIHAN": ["12-16 Jan 2026"], "TEMPAT": ["Pusdiklat"], "NO": [1], "NAMA PEGAWAI": ["Fajar"], "NIP": ["199901012024121001"], "PANGKAT": ["II/c"], "SATUAN KERJA": ["KPU Batam"]})
            buffer = io.BytesIO(); 
            with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer: df_dummy.to_excel(writer, index=False)
            buffer.seek(0)
            st.download_button("📥 Template Peserta", buffer, "Template_Peserta.xlsx", use_container_width=True)
        with sc2:
            if st.button("🔄 Reset", type="secondary", use_container_width=True): reset_app()

    with c_ttd:
        st.markdown("###### 2. Pejabat Tanda Tangan")
        nama_ttd = st.text_input("Nama Pejabat", "Ayu Sukorini")
        jabatan_ttd = st.text_input("Jabatan", "Sekretaris Direktorat Jenderal")

    with c_nd:
        st.markdown("###### 3. Detail Nota Dinas")
        nomor_nd = st.text_input("Nomor ND", "[@NomorND]")
        tanggal_nd = st.text_input("Tanggal ND", "[@TanggalND]")

    st.divider()

    if uploaded_file:
        try:
            df_raw = pd.read_excel(uploaded_file, dtype=str)
            clean_cols = {}
            for col in df_raw.columns:
                upper_col = col.strip().upper().replace(" ", "_").replace("-", "_")
                if "NAMA" in upper_col: clean_cols[col] = "NAMA"
                elif "NIP" in upper_col: clean_cols[col] = "NIP"
                elif "PANGKAT" in upper_col or "GOL" in upper_col: clean_cols[col] = "PANGKAT"
                elif "KERJA" in upper_col or "SATKER" in upper_col: clean_cols[col] = "SATKER"
                elif "TEMPAT" in upper_col: clean_cols[col] = "TEMPAT"
                elif "JUDUL" in upper_col or "DIKLAT" in upper_col: clean_cols[col] = "JUDUL_PELATIHAN"
                elif "TANGGAL" in upper_col: clean_cols[col] = "TANGGAL_PELATIHAN"
                else: clean_cols[col] = upper_col 
            df_raw = df_raw.rename(columns=clean_cols).fillna("-")
            
            # Auto-Detect NIP
            if 'NIP' in df_raw.columns:
                df_raw['USIA'] = df_raw['NIP'].apply(calculate_age_from_nip)
                df_raw['GENDER'] = df_raw['NIP'].apply(get_gender_from_nip)
            else:
                df_raw['USIA'] = None; df_raw['GENDER'] = "Tidak Diketahui"

            # Auto-Detect Title from Calendar
            sh = connect_to_gsheet()
            if sh and 'JUDUL_PELATIHAN' in df_raw.columns:
                try:
                    ws = sh.worksheet(SHEET_KALENDER)
                    data_cal = pd.DataFrame(ws.get_all_records())
                    judul_peserta = df_raw['JUDUL_PELATIHAN'].iloc[0]
                    if not data_cal.empty and judul_peserta in data_cal['JUDUL_PELATIHAN'].values:
                        st.success(f"✅ Pelatihan '{judul_peserta}' terdaftar di Kalender. Status akan diupdate setelah download.")
                except: pass

            st.markdown("###### 4. Preview & Edit Data")
            df_edited = st.data_editor(df_raw, num_rows="dynamic", use_container_width=True)
            
            ts = datetime.datetime.now().strftime("%H%M%S")
            word_buffer = generate_word_combined(df_edited, nama_ttd, jabatan_ttd, nomor_nd, tanggal_nd)
            zip_buffer = generate_zip_files(df_edited, nama_ttd, jabatan_ttd, nomor_nd, tanggal_nd)
            
            st.markdown("<br>", unsafe_allow_html=True)
            c_d1, c_d2 = st.columns(2)
            with c_d1: 
                st.download_button("📄 Download Lampiran ND (.docx)", word_buffer, f"Lampiran_{ts}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary", use_container_width=True, on_click=save_to_cloud_callback, args=(df_edited,))
            with c_d2: 
                st.download_button("📦 Download Arsip ZIP", zip_buffer, f"Arsip_{ts}.zip", "application/zip", use_container_width=True, on_click=save_to_cloud_callback, args=(df_edited,))

        except Exception as e: st.error(f"Error: {e}")
    else: st.info("👈 Silakan upload file Excel peserta.")

# --- TAB DASHBOARD ---
with tab_dash:
    if uploaded_file and 'df_edited' in locals():
        df_viz = df_edited
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Total Peserta", len(df_viz))
        jml_pelatihan = df_viz['JUDUL_PELATIHAN'].nunique() if 'JUDUL_PELATIHAN' in df_viz.columns else 0
        c2.metric("Jumlah Pelatihan", jml_pelatihan)
        avg_age = df_viz['USIA'].mean() if 'USIA' in df_viz.columns else 0
        c3.metric("Rata-rata Usia", f"{avg_age:.0f} Tahun" if pd.notna(avg_age) else "-")
        c4.metric("Satker", df_viz['SATKER'].nunique() if 'SATKER' in df_viz.columns else 0)
        st.markdown("---")
        col_g1, col_g2 = st.columns(2)
        with col_g1:
            if 'USIA' in df_viz.columns and df_viz['USIA'].notna().any():
                age_counts = df_viz['USIA'].dropna().astype(int).value_counts().sort_index()
                st.bar_chart(age_counts, color="#3498DB")
        with col_g2:
            if 'GENDER' in df_viz.columns:
                gender_counts = df_viz['GENDER'].value_counts()
                fig, ax = plt.subplots(figsize=(5, 4))
                ax.pie(gender_counts, labels=gender_counts.index, autopct='%1.1f%%', startangle=90, colors=['#3498DB', '#E91E63', '#95A5A6'])
                st.pyplot(fig); plt.close(fig)
    
    st.markdown("#### 📅 Monitoring Realisasi Diklat 2026")
    sh = connect_to_gsheet()
    if sh:
        try:
            ws = sh.worksheet(SHEET_KALENDER)
            df_cal = pd.DataFrame(ws.get_all_records())
            if not df_cal.empty:
                total_plan = len(df_cal)
                total_done = len(df_cal[df_cal['STATUS'] == 'Selesai'])
                progress = total_done / total_plan if total_plan > 0 else 0
                st.progress(progress, text=f"Realisasi: {total_done} dari {total_plan} Pelatihan ({progress:.1%})")
                st.dataframe(df_cal[['JUDUL_PELATIHAN', 'RENCANA_TANGGAL', 'LOKASI', 'STATUS', 'REALISASI']], use_container_width=True)
        except: st.info("Data Kalender belum tersedia.")

# --- TAB DATABASE (DANGER ZONE) ---
with tab_db:
    st.subheader("🔗 Log Peserta")
    sh = connect_to_gsheet()
    if sh:
        ws = sh.worksheet(SHEET_HISTORY)
        st.dataframe(pd.DataFrame(ws.get_all_records()))
        
        st.markdown("---")
        with st.expander("⚠️ DANGER ZONE / AREA BERBAHAYA"):
            st.markdown("""<div class="danger-box">
            <b>PERINGATAN:</b> Tindakan di bawah ini bersifat destruktif dan tidak dapat dibatalkan.
            Harap berhati-hati sebelum menekan tombol reset.
            </div>""", unsafe_allow_html=True)
            
            c_d1, c_d2 = st.columns(2)
            
            # FITUR 1: RESET STATUS KALENDER
            with c_d1:
                st.markdown("##### 1. Reset Status Kalender")
                st.caption("Mengembalikan semua status pelatihan menjadi 'Pending'.")
                confirm_cal = st.checkbox("Saya sadar ini akan mereset progress.", key="chk_cal")
                if confirm_cal:
                    if st.button("🔴 RESET STATUS KALENDER", type="primary"):
                        reset_calendar_status()
            
            # FITUR 2: HAPUS LOG PESERTA
            with c_d2:
                st.markdown("##### 2. Hapus Log Peserta")
                st.caption("Menghapus semua riwayat upload peserta (Sheet1).")
                confirm_log = st.checkbox("Saya sadar data akan hilang permanen.", key="chk_log")
                if confirm_log:
                    if st.button("🔴 HAPUS SEMUA LOG", type="primary"):
                        clear_history_log()
